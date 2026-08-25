"""Generate Word Technical Specification documents from the CPI graph."""

from __future__ import annotations

from collections import defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from src.agent.tools import (
    get_adapter_security_summary,
    get_error_handling_coverage,
    get_iflow_processes,
    get_subflow_chain,
)
from src.docgen.ts_document import build_ts_content
from src.graph.neo4j_client import run_query


TABLE_STYLE = "Light Grid Accent 1"


def _set_table_style(table: Any) -> None:
    """Apply an available built-in style across Word installations."""
    try:
        table.style = TABLE_STYLE
    except KeyError:
        table.style = "Table Grid"


def _add_table(document: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    _set_table_style(table)
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = "" if value is None else str(value)


def _process_flow(artifact_id: str) -> list[dict[str, Any]]:
    """Return a deterministic depth-first ordering from process start steps."""
    query = """
    MATCH (i:IFlow {id: $artifact_id})<-[:PART_OF]-(p:Process)<-[:BELONGS_TO]-(s:Step)
    OPTIONAL MATCH (s)-[n:NEXT]->(target:Step)
    OPTIONAL MATCH (s)-[:USES]->(r:Resource)
    RETURN s.step_key AS step_key, s.id AS step_id, s.name AS step_name,
           s.activity_type AS activity_type, s.bpmn_type AS bpmn_type,
           p.classification AS classification,
           collect(DISTINCT {target_key: target.step_key, target_name: target.name,
                             condition: n.condition}) AS next_steps,
           collect(DISTINCT r.filename) AS resources
    """
    rows = run_query(query, {"artifact_id": artifact_id})
    if not rows:
        return []

    steps = {
        row["step_key"]: {
            "id": row["step_id"],
            "name": row["step_name"],
            "type": row["activity_type"] or row["bpmn_type"] or "",
            "classification": row["classification"],
            "resources": [resource for resource in row["resources"] if resource],
            "next_steps": [item for item in row["next_steps"] if item["target_key"]],
        }
        for row in rows
    }
    incoming = {
        target["target_key"]
        for step in steps.values()
        for target in step["next_steps"]
    }
    classifications = {"main": 0, "local_subprocess": 1, "error_handling": 2}
    starts = sorted(
        (key for key in steps if key not in incoming),
        key=lambda key: (classifications.get(steps[key]["classification"], 9), steps[key]["id"] or ""),
    )
    ordered: list[dict[str, Any]] = []
    visited: set[str] = set()

    def visit(step_key: str) -> None:
        if step_key in visited or step_key not in steps:
            return
        visited.add(step_key)
        ordered.append(steps[step_key])
        for next_step in sorted(steps[step_key]["next_steps"], key=lambda item: item["target_key"]):
            visit(next_step["target_key"])

    for start in starts:
        visit(start)
    for step_key in sorted(steps, key=lambda key: (classifications.get(steps[key]["classification"], 9), steps[key]["id"] or "")):
        visit(step_key)
    return ordered


def _error_process_details(artifact_id: str) -> dict[str, dict[str, Any]]:
    coverage = get_error_handling_coverage()
    if "error" in coverage:
        return {}
    for iflow in coverage["iflows_with_error_handling"]:
        if iflow["artifact_id"] == artifact_id:
            return {detail["process_id"]: detail for detail in iflow["details"]}
    return {}


def generate_ts_docx(artifact_id: str) -> str:
    """Generate a Word Technical Specification for an iFlow and return its path."""
    content = build_ts_content(artifact_id)
    if not content.get("artifact_id"):
        raise ValueError(f"IFlow not found: {artifact_id}")

    document = Document()
    document.add_heading(f"{content['artifact_id']} (v{content['version']})", level=1)

    document.add_heading("Scope", level=2)
    document.add_paragraph(content["scope_summary"])

    document.add_heading("Process Flow", level=2)
    for index, step in enumerate(_process_flow(artifact_id), start=1):
        resources = f" — resources: {', '.join(step['resources'])}" if step["resources"] else ""
        paragraph = document.add_paragraph(
            f"Step {index}: {step['name']} ({step['type']}){resources}",
            style="List Number",
        )
        for next_step in step["next_steps"]:
            if next_step["condition"]:
                document.add_paragraph(
                    f"→ if {next_step['condition']}: goes to {next_step['target_name']}",
                    style="List Bullet 2",
                )

    processes = get_iflow_processes(artifact_id)
    non_main_processes = [
        process for process in processes.get("processes", [])
        if process["classification"] != "main"
    ]
    if non_main_processes:
        document.add_heading("Process Structure", level=2)
        error_details = _error_process_details(artifact_id)
        for process in non_main_processes:
            label = process["classification"].replace("_", " ")
            line = f"{process['process_id']} — {label}, {process['step_count']} steps"
            if process["classification"] == "error_handling":
                triggers = error_details.get(process["process_id"], {}).get("trigger_types", [])
                if triggers:
                    line += f"; triggers: {', '.join(triggers)}"
            document.add_paragraph(line, style="List Bullet")

    mapping_complexity = {
        resource.get("filename"): resource.get("has_complex_transformations")
        for resource in content["resources"]
    }
    if content["field_mappings"]:
        document.add_heading("Field Mappings", level=2)
        rows = []
        for mapping in content["field_mappings"]:
            function = mapping.get("function") or ""
            if mapping_complexity.get(mapping.get("resource")):
                function = f"{function} (complex transformation)".strip()
            rows.append([mapping.get("source"), mapping.get("target"), function])
        _add_table(document, ["Source", "Target", "Function"], rows)

    document.add_heading("Systems & Endpoints", level=2)
    security = get_adapter_security_summary(artifact_id)
    security_by_step_system = {
        (adapter["step_name"], adapter["system_name"]): adapter
        for adapter in security.get("adapters", [])
    }
    system_rows = []
    for system in content["systems"]:
        adapter = security_by_step_system.get((system.get("step_name"), system.get("system_name")), {})
        # get_iflow_systems omits step_name, so match the unique system fallback.
        if not adapter:
            adapter = next(
                (item for item in security.get("adapters", []) if item["system_name"] == system.get("system_name")),
                {},
            )
        counts = ""
        if adapter:
            counts = f"{adapter.get('externalized_count', 0)} / {adapter.get('literal_count', 0)}"
        system_rows.append([
            system.get("system_name"), system.get("direction"), system.get("component_type"),
            system.get("address"), counts,
        ])
    _add_table(
        document,
        ["System", "Direction", "Protocol", "Address", "Externalized/Literal Properties"],
        system_rows,
    )

    document.add_heading("Externalized Parameters", level=2)
    for category, parameters in content["parameters"].items():
        document.add_heading(category, level=3)
        rows = [
            [parameter.get("name"), parameter.get("label"), parameter.get("value"), parameter.get("is_required")]
            for parameter in parameters
        ]
        _add_table(document, ["Name", "Label", "Value", "Required"], rows)

    subflows = get_subflow_chain(artifact_id)
    calls = subflows.get("calls", [])
    called_by = subflows.get("called_by", [])
    if calls or called_by:
        document.add_heading("Subflow Relationships", level=2)
        for call in calls:
            document.add_paragraph(
                f"Calls: {call['callee']} (via {call['address']})", style="List Bullet"
            )
        for caller in called_by:
            document.add_paragraph(
                f"Called by: {caller['caller']} (via {caller['address']})", style="List Bullet"
            )

    output_dir = Path("output/ts_documents")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{artifact_id}.docx"
    document.save(output_path)
    return str(output_path)
